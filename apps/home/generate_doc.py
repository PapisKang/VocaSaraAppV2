import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from PIL import Image as PILIMAGE
from io import BytesIO
from datetime import date
from flask import Blueprint, flash, session
from flask_login import current_user
from datetime import datetime
import math
import json
from apps.authentication.models import UserProfile, Users, ImageUploadVisible, ImageUploadInvisible, RapportGenere, DocumentRapportGenere
from apps import db
from flask import render_template, Flask, redirect, request, url_for, jsonify, send_file
from sqlalchemy import func
from sqlalchemy import select
import logging



@Blueprint.route('/generate_report_document', methods=['GET', 'POST'])
def generate_report_document():
    success_message = None
    error_message = None
    rapports = []
    try:
        current_date = datetime.now().strftime("%Y-%m-%d")
        last_image_data = ImageUploadVisible.query.order_by(
            ImageUploadVisible.upload_date.desc()).first()

        if not last_image_data:
            error_message = "Aucune donnée disponible dans la base de données."
            return render_template('rapport/creer_un_rapport.html')

        date = datetime.now().strftime("%Y-%m-%d")
        nom_operateur = last_image_data.nom_operateur
        feeder = last_image_data.feeder
        zone = last_image_data.zone
        groupement = last_image_data.groupement_troncon
        image_data = ImageUploadVisible.query.filter_by(
            nom_operateur=nom_operateur).all()

        # Charger les normes_conseils des défauts à partir du fichier JSON
        with open('.apps/phrase_normes_conseils/normes_conseils.json', encoding='utf-8') as f:
            normes_conseils_data = json.load(f)
        # Fonction pour créer une table des matières automatique

        def create_table_of_contents(document):
            # ////// Insérer la table des matières
            table_of_contents = document.add_paragraph("Table des matières")
            table_of_contents.runs[0].font.size = Pt(16)
            table_of_contents.runs[0].font.name = "Times New Roman"
            table_of_contents.runs[0].underline = True
            table_of_contents.runs[0].font.color.rgb = RGBColor(
                0x2F, 0x54, 0x96)  # Bleu sombre (#2F5496)
            document.add_paragraph(
                f"RAPPORT D’INSPECTION PAR DRONE DANS LA ZONE {zone}")
            feeder_title = document.add_paragraph()
            run = feeder_title.add_run("FEEDER : ")
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"
            document.add_paragraph(f"\t{feeder}")
            groupement_title = document.add_paragraph()
            run = groupement_title.add_run("GROUPEMENT : ")
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"
            document.add_paragraph(f"\tGROUPEMENT TRONCONS ENTRE {groupement}")
            return document
        # Créer le document Word
        document = Document()
        # Paramètres de mise en page
        sections = document.sections
        for section in sections:
            section.top_margin = Pt(52)  # Marge supérieure
            section.bottom_margin = Pt(52)  # Marge inférieure
            section.left_margin = Pt(52)  # Marge gauche
            section.right_margin = Pt(52)  # Marge droite
            section.page_width = Pt(612)  # Largeur de page
            section.page_height = Pt(792)  # Hauteur de page

        # Ajouter la première page
        first_page = document.add_paragraph()
        first_page.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = first_page.add_run(
            f"RAPPORT D’INSPECTION PAR DRONE DANS LA ZONE {zone}")
        run.bold = True
        run.font.size = Pt(28)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(
            0x2F, 0x54, 0x96)  # Bleu sombre (#2F5496)

        # Créer la table des matières
        document = create_table_of_contents(document)
        # Ajouter la distance entre l'entête et la page
        for section in sections:
            section.header_distance = Pt(36)

        # Ajouter la deuxième page
        document.add_page_break()
        second_page = document.add_paragraph()
        second_page.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = second_page.add_run("FEEDER:")
        run.bold = True
        run.underline = True
        run.font.size = Pt(20)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(
            0x2F, 0x54, 0x96)  # Bleu sombre (#2F5496)
        second_page.add_run(f"\n\n\t{feeder}\n\n")
        run = second_page.add_run("GROUPEMENT:")
        run.bold = True
        run.underline = True
        run.font.size = Pt(20)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(
            0x2F, 0x54, 0x96)  # Bleu sombre (#2F5496)
        second_page.add_run(f"\n\n\tGROUPEMENT TRONCONS ENTRE {groupement}")
        run.font.size = Pt(20)
        run.font.name = "Times New Roman"
        # Ajouter la troisième page
        document.add_page_break()
        third_page = document.add_paragraph()
        third_page.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = third_page.add_run("\n")
        run.add_picture("{{ config.ASSETS_ROOT }}/img/vocasara/logo.png", width=Inches(7))
        run = third_page.add_run("\n\n")
        run = third_page.add_run(f"GROUPEMENT TRONCONS ENTRE {groupement}")
        run.underline = True
        run.font.size = Pt(16)
        for image_info in image_data:
            # Charger l'image à partir des données binaires
            image = PILIMAGE.open(BytesIO(image_info.data))

            # Récupérer les informations sur les défauts
            # Supposant que les défauts sont séparés par des barres obliques
            defects = image_info.type_defaut.split('/')

            for defect in defects:
                defect = defect.strip()
                # Accédez aux remarques et conseils depuis le fichier JSON
                remarque = "Remarque par défaut"
                conseil = "Conseil par défaut"
                for defaut, info in normes_conseils_data.items():
                    if defect in defaut:
                        remarque = info.get("I", "Remarque par défaut")
                        conseil = info.get("J", "Conseil par défaut")
                        break

                # Ajouter une page pour chaque défaut
                document.add_page_break()
                page = document.add_paragraph()
                page.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                # Créer le tableau pour stocker les données de l'image
                table = document.add_table(rows=5, cols=1)
                table.style = "Table Grid"
                table.columns[0].width = Pt(900)

                # Cellule pour les défauts
                cell = table.cell(0, 0)
                cell.text = f"Défauts : {defect}"
                cell.paragraphs[0].runs[0].bold = True

                # Reduce image dimensions
                max_image_size = (800, 600)
                try:
                    image.thumbnail(max_image_size, PILIMAGE.ANTIALIAS)
                except:
                    image.thumbnail(max_image_size)

                # Compress image to reduce size
                image = image.convert("RGB")
                image_bytes = BytesIO()
                image.save(image_bytes, format='JPEG', quality=95)
                image_bytes.seek(0)

                # Cellule pour l'image
                cell = table.cell(1, 0)
                cell.vertical_alignment = 1  # Alignement vertical centré
                cell.add_paragraph().add_run().add_picture(image_bytes, width=Inches(7))

                # Cellule pour la localisation
                cell = table.cell(2, 0)
                cell.vertical_alignment = 1  # Alignement vertical centré
                # Cellule pour la remarque
                cell = table.cell(3, 0)
                cell.vertical_alignment = 1
                cell.add_paragraph().add_run("Remarque: ").bold = True
                cell.add_paragraph(remarque)

                cell = table.cell(4, 0)
                cell.vertical_alignment = 1
                cell.add_paragraph().add_run("Conseil: ").bold = True
                cell.add_paragraph(conseil)

        # Ajouter le pied de page
        footer_text = f"RAPPORT DU {date}\t\t{nom_operateur}"
        sections[-1].footer.paragraphs[0].text = footer_text
        # En-tête
        header_text = "VOCASARA S.U.A.R.L"
        header_paragraph = sections[0].header.paragraphs[0]
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_run = header_paragraph.add_run(header_text)
        header_run.bold = True
        header_run.font.size = Pt(22)
        header_run.font.name = "Times New Roman"
        header_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Noir
        header_run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Logo
        logo_path = "{{config.ASSETS_ROOT}}/img/vocasara/logo.png"
        header_paragraph.add_run().add_picture(logo_path, width=Inches(4)
                                               ).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Sauvegarder le document Word
        # Convert the date to a string with the format "YYYY-MM-DD"
        date = datetime.now().strftime("%Y-%m-%d")
        # document.save("rapport_du_feeder_{}_du_{}.docx".format(feeder,date))
        nom_du_fichier = ("Rapport_defauts_visible_du_feeder_{}_du_{}_par_{}.docx".format(
            feeder, current_date, nom_operateur))
        # Enregistrer le fichier Word en mémoire
        word_file = io.BytesIO()
        document.save(word_file)
        word_file.seek(0)
        # Enregistrer le fichier dans la base de données
        document_word = DocumentRapportGenere(
            nom_operateur=nom_operateur,
            nom_du_rapport=nom_du_fichier,
            data=word_file.read()
        )
        db.session.add(document_word)
        db.session.commit()

        success_message = "Le rapport a été généré avec succès."
        rapports = RapportGenere.query.all()
    except FileNotFoundError:
        error_message = "Le fichier de template n'a pas été trouvé."
    except Exception as e:
        logging.error(
            f"Une erreur s'est produite dans genernate doc: {str(e)}")
        error_message = f"Une erreur s'est produite : {str(e)}"

    return render_template('rapport/creer_un_rapport.html', rapports=rapports, success_message=success_message, error_message=error_message)
