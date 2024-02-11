# -*- encoding: utf-8 -*-
from apps.home import blueprint
from flask import render_template, Flask, redirect, request, url_for
from flask_login import login_required
from jinja2 import TemplateNotFound
from apps import db
from flask_wtf import FlaskForm
from wtforms import StringField, DateField, SelectField
from wtforms.validators import DataRequired
from datetime import datetime
from apps.authentication.models import Troncon, Feeder, RapportGenere
from flask_login import current_user
from werkzeug.utils import secure_filename
from io import BytesIO
import zlib
import base64
import logging
import os
import base64
from apps.authentication.models import UserProfile, Users, ImageUploadVisible, ImageUploadInvisible, RapportGenere, DocumentRapportGenere
from flask import render_template, jsonify, send_file
import json
from sqlalchemy import func
from sqlalchemy import select

import xlsxwriter
import openpyxl
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, Border, Side
from openpyxl.chart import DoughnutChart, Reference
from openpyxl.styles import Font, Alignment
from openpyxl.utils.dataframe import dataframe_to_rows

import io
import logging
from apps.config import Config



from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image as PILIMAGE

# Ajouter un log lorsqu'un utilisateur se connecte
@blueprint.before_request
def log_request():
    if current_user.is_authenticated:
        access_logger.info('Utilisateur connecté : %s | Accès à l\'URL : %s', current_user.username, request.url)
    else:
        access_logger.info('Accès à l\'URL : %s', request.url)

@blueprint.route('/<template>')
@login_required
def route_template(template):
    print(template)
    try:
        if not template.endswith('.html'):
            template += '.html'

        # Detect the current page
        segment, active_menu, parent, segment_name = get_segment(request)

        # Serve the file (if exists) from app/templates/home/FILE.html
        return render_template(
            "home/" + template,
            segment=segment,
            active_menu=active_menu,
            parent=parent,
            segment_name=segment_name
        )

    except TemplateNotFound:
        return render_template('home/page-404.html'), 404

    except:
        return render_template('home/page-500.html'), 500


# Helper - Extract current page name from request
def get_segment(request):

    try:

        segment = request.path.split('/')[-1]
        active_menu = None
        parent = ""
        segment_name = ""

        core = segment.split('.')[0]

        if segment == '':
            segment = 'index'

        parent = core.split('-')[0] if core.split('-')[0] else ""
        segment_name = core

        return segment, active_menu, parent, segment_name

    except:
        return 'index', ''


# Route pour afficher le formulaire et les listes déroulantes
@blueprint.route('/generation_rapport', methods=['GET'])
def generation_rapport_form():
    feeders_existants = Feeder.query.all()
    troncons_existants = Troncon.query.all()
    return render_template('rapport/generation_rapport.html', feeders_existants=feeders_existants, troncons_existants=troncons_existants)


@blueprint.route('/generate_rapport', methods=['POST'])
@login_required
def generate_rapport():
    feeders_existants = Feeder.query.all()
    troncons_existants = Troncon.query.all()

    feeder_nom = request.form.get('feeder') or request.form.get('feederInput')
    troncon_nom = request.form.get(
        'troncon') or request.form.get('tronconInput')
    date_debut_str = request.form.get('dateDebut')
    date_fin_str = request.form.get('dateFin')
    operateur = current_user.email
    groupement_troncon = request.form.get('groupementTroncon')
    zone = request.form.get('zone')

    # Vérifier si les champs obligatoires sont vides
    if not feeder_nom or not troncon_nom or not date_debut_str or not date_fin_str or not groupement_troncon or not zone:
        return render_template('rapport/generation_rapport.html', feeders_existants=feeders_existants, troncons_existants=troncons_existants, error_message='Veuillez remplir tous les champs obligatoires.')

    # Vérifier si les champs de date sont vides
    if not date_debut_str or not date_fin_str:
        return render_template('rapport/generation_rapport.html', feeders_existants=feeders_existants, troncons_existants=troncons_existants, error_message='Veuillez remplir les champs de date.')

    try:
        date_debut = datetime.strptime(date_debut_str, '%Y-%m-%d')
        date_fin = datetime.strptime(date_fin_str, '%Y-%m-%d')
    except ValueError:
        return render_template('rapport/generation_rapport.html', feeders_existants=feeders_existants, troncons_existants=troncons_existants, error_message='Format de date incorrect.')

    existing_feeder = Feeder.query.filter_by(Nom=feeder_nom).first()
    if existing_feeder:
        feeder_id = existing_feeder.id
    else:
        new_feeder = Feeder(Nom=feeder_nom)
        db.session.add(new_feeder)
        db.session.commit()
        feeder_id = new_feeder.id

    existing_troncon = Troncon.query.filter_by(Nom=troncon_nom).first()
    if existing_troncon:
        troncon_id = existing_troncon.id
    else:
        new_troncon = Troncon(Nom=troncon_nom)
        db.session.add(new_troncon)
        db.session.commit()
        troncon_id = new_troncon.id

    new_report = RapportGenere(
        nom_operateur=operateur,
        feeder=feeder_nom,
        troncon=troncon_nom,
        date_debut=date_debut,
        date_fin=date_fin,
        zone=zone,
        groupement_troncon=groupement_troncon
    )
    db.session.add(new_report)
    db.session.commit()

    # Ajouter l'id du rapport généré aux cookies
    response = redirect(url_for('home_blueprint.confirmation_page'))
    response.set_cookie('rapportGenereId', str(new_report.id))

    return response

# Nouvelle route pour afficher la page avec les deux boutons et le texte explicatif

@login_required
@blueprint.route('/confirmation_page', methods=['GET'])
def confirmation_page():
    return render_template('rapport/confirmation_page.html')


@blueprint.route('/apropos')
@login_required
def apropos():
    return render_template('home/apropos.html', segment='apropos')


@blueprint.route('/acceuil')
@login_required
def acceuil():
    return render_template('home/acceuil.html')


#./////////////////////// page de localisation de défuat§§§§§§§§§§
@blueprint.route('/update_status/<int:image_id>', methods=['POST'])
@login_required
def update_status(image_id):
    # Récupérer l'utilisateur connecté
    user = current_user
    # Récupérer le statut envoyé dans la requête
    new_status = request.json.get('new_status')
    # Récupérer l'image à mettre à jour
    image = ImageUploadVisible.query.get(image_id)
    if not image:
        return jsonify({'error': 'Image not found'}), 404
    # Mettre à jour les informations
    image.status = new_status
    image.updated_by = user.username  # Mettre à jour avec le nom d'utilisateur de l'utilisateur connecté
    image.update_date = datetime.utcnow()  # Mettre à jour avec la date actuelle
    # Sauvegarder les modifications dans la base de données
    db.session.commit()
    return jsonify({'success': 'Status updated successfully'}), 200


@blueprint.route('/get_rapports')
def get_rapports():
    rapports = RapportGenere.query.all()
    rapports_data = [
        {
            'id': rapport.id,
            'nom_operateur': rapport.nom_operateur,
            'feeder': rapport.feeder,
            'troncon': rapport.troncon,
            'date_debut': rapport.date_debut.strftime('%Y-%m-%d %H:%M:%S'),
            'date_fin': rapport.date_fin.strftime('%Y-%m-%d %H:%M:%S'),
            'zone': rapport.zone,
            'date_created': rapport.date_created.strftime('%Y-%m-%d %H:%M:%S')
        }
        for rapport in rapports
    ]
    return jsonify(rapports_data)

@blueprint.route('/get_default_types')
def get_default_types():
    # Query distinct default types from the database
    default_types = db.session.query(ImageUploadVisible.type_defaut).distinct().all()
    default_types = [row[0] for row in default_types if row[0] is not None]
    return jsonify(default_types)


@blueprint.route('/get_map_data')
def get_map_data():
    page = request.args.get('page', default=1, type=int)
    per_page = request.args.get('per_page', default=20, type=int)

    # Récupérer l'utilisateur connecté
    user = current_user
    # Get the selected default type
    selected_default_type = request.args.get('default_type')

    # Query the data based on the selected default type
    if selected_default_type:
        image_points = ImageUploadVisible.query.filter_by(type_defaut=selected_default_type)
    else:
        image_points = ImageUploadVisible.query
    # Récupérer le statut sélectionné (par défaut "en attente")
    selected_status = request.args.get('status', default='en attente')

    # Récupérer l'ID du rapport généré sélectionné
    rapport_id = request.args.get('rapport_id')

    # Initialiser image_points en fonction de la présence ou non de rapport_id
    if rapport_id:
        image_points = ImageUploadVisible.query.filter_by(rapport_genere_id=rapport_id)
    else:
        image_points = ImageUploadVisible.query

    # Vérifier si l'utilisateur est un administrateur
    if user.role == Config.USERS_ROLES['ADMIN']:
        # Si l'utilisateur est un administrateur, récupérer toutes les données
        image_points = image_points.filter(
            ImageUploadVisible.type_defaut.isnot(None),
            ImageUploadVisible.type_defaut != "",
            ImageUploadVisible.status.isnot(None)  # Ajouter cette condition pour filtrer les points avec un statut non nul
        ).paginate(page=page, per_page=per_page, error_out=False)
    else:
        # Sinon, récupérer seulement les données de l'utilisateur connecté
        image_points = image_points.filter(
            ImageUploadVisible.type_defaut.isnot(None),
            ImageUploadVisible.type_defaut != "",
            ImageUploadVisible.status.isnot(None)  # Ajouter cette condition pour filtrer les points avec un statut non nul
            # Ajouter d'autres conditions si nécessaire pour filtrer par utilisateur
        ).paginate(page=page, per_page=per_page, error_out=False)

    map_data = []
    for point in image_points.items:
        # Filtrer les points en fonction du statut sélectionné
        if point.status == selected_status:
            data = {
                'imageId': point.id,
                'latitude': point.latitude,
                'longitude': point.longitude,
                'type_defaut': point.type_defaut,
                'feeder': point.feeder,
                'troncon': point.troncon,
                'zone': point.zone,
                'filename': point.filename,
                'nom_operateur': point.nom_operateur,
                'upload_date': point.upload_date.strftime('%Y-%m-%d %H:%M:%S'),
                'image_binary': point.data,
                'status': point.status  
            }
            map_data.append(data)

    return jsonify(map_data)



@blueprint.route('/localisation_page')
@login_required
def localisation_page():
    return render_template('home/localisation_defauts.html')

#./////////////////////// page de localisation de défuat§§§§§§§§§§

#//////////////////////page de statistique.////////////////

@blueprint.route('/statistiques')
@login_required
def statistiques():
    # Récupérer les statistiques des types de défauts avec les informations supplémentaires
    type_defaut_stats = db.session.query(
        ImageUploadVisible.type_defaut,
        db.func.count().label('defaut_count')
    ).filter(
        ImageUploadVisible.type_defaut.isnot(None)  # Exclude rows where type_defaut is null
    ).group_by(
        ImageUploadVisible.type_defaut
    ).all()

    # Récupérer les dates
    dates = db.session.query(ImageUploadVisible.upload_date).all()

    # Convertir les données en format approprié pour les graphiques
    labels = [row[0] for row in type_defaut_stats]
    values = [row.defaut_count for row in type_defaut_stats]

    # Préparer les données pour le frontend
    data = {
        'labels': labels,
        'values': values,
        'dates': dates,
    }

    # Convertir les données en JSON
    data_json = json.dumps(data, default=str)

    # Rendre la page HTML avec les données
    return render_template('statistics/statistics.html', data=data_json)



#//////////////////////page de statistique.////////////////


@blueprint.route('/localisation_defauts_invisible_page')
@login_required
def localisation_defauts_invisible_page():
    return render_template('/home/localisation_defauts_invisible.html')

 
@blueprint.route('/statistics_invisible')
def statistics_invisible():
    return render_template('/statistics/statistics_invisible.html')

#./////////////////////////Partie inspections.///////////////


@blueprint.route('/rapport_id_page')
def rapport_id_page():
    rapports = RapportGenere.query.all()
    return render_template('rapport/rapport_id_page.html',  rapports=rapports)

@login_required
@blueprint.route('/mes_inspections/<int:rapport_id>', methods=['GET'])
def mes_inspections(rapport_id):
    rapport = RapportGenere.query.get(rapport_id)
    if rapport:
        images_visibles = ImageUploadVisible.query.filter_by(
            rapport_genere_id=rapport_id).all()
        return render_template('rapport/mes_inspections.html', rapport=rapport, images_visibles=images_visibles)
    else:
        return redirect(url_for('authentication_blueprint.index'))


@blueprint.route('/generate_report_document_page')
def generate_report_document_page():
    rapports = RapportGenere.query.all()
    return render_template('rapport/creer_un_rapport.html', rapports=rapports)


@blueprint.route('/generate_report_document', methods=['GET', 'POST'])
def generate_report_document():
    success_message = None
    error_message = None
    rapports = []
    try:
        current_date = datetime.now().strftime("%Y-%m-%d")
        last_image_data = ImageUploadVisible.query.order_by(
            ImageUploadVisible.upload_date.desc()).first()

        if not last_image_data:
            error_message = "Aucune donnée disponible dans la base de données."
            return render_template('rapport/creer_un_rapport.html')

        date = datetime.now().strftime("%Y-%m-%d")
        nom_operateur = last_image_data.nom_operateur
        feeder = last_image_data.feeder
        zone = last_image_data.zone
        groupement = last_image_data.groupement_troncon
        image_data = ImageUploadVisible.query.filter_by(
            nom_operateur=nom_operateur).all()

        # Charger les normes_conseils des défauts à partir du fichier JSON
        with open('./apps/phrase_normes_conseils/normes_conseils.json', encoding='utf-8') as f:
            normes_conseils_data = json.load(f)

        # Fonction pour créer une table des matières automatique

        def create_table_of_contents(document):
            # ////// Insérer la table des matières
            table_of_contents = document.add_paragraph("Table des matières")
            table_of_contents.runs[0].font.size = Pt(16)
            table_of_contents.runs[0].font.name = "Times New Roman"
            table_of_contents.runs[0].underline = True
            table_of_contents.runs[0].font.color.rgb = RGBColor(
                0x2F, 0x54, 0x96)  # Bleu sombre (#2F5496)
            document.add_paragraph(
                f"RAPPORT D’INSPECTION PAR DRONE DANS LA ZONE {zone}")
            feeder_title = document.add_paragraph()
            run = feeder_title.add_run("FEEDER : ")
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"
            document.add_paragraph(f"\t{feeder}")
            groupement_title = document.add_paragraph()
            run = groupement_title.add_run("GROUPEMENT : ")
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"
            document.add_paragraph(f"\tGROUPEMENT TRONCONS ENTRE {groupement}")
            return document
        # Créer le document Word
        document = Document()
        # Paramètres de mise en page
        sections = document.sections
        for section in sections:
            section.top_margin = Pt(52)  # Marge supérieure
            section.bottom_margin = Pt(52)  # Marge inférieure
            section.left_margin = Pt(52)  # Marge gauche
            section.right_margin = Pt(52)  # Marge droite
            section.page_width = Pt(612)  # Largeur de page
            section.page_height = Pt(792)  # Hauteur de page

        # Ajouter la première page
        first_page = document.add_paragraph()
        first_page.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = first_page.add_run(
            f"RAPPORT D’INSPECTION PAR DRONE DANS LA ZONE {zone}")
        run.bold = True
        run.font.size = Pt(28)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(
            0x2F, 0x54, 0x96)  # Bleu sombre (#2F5496)

        # Créer la table des matières
        document = create_table_of_contents(document)
        # Ajouter la distance entre l'entête et la page
        for section in sections:
            section.header_distance = Pt(36)

        # Ajouter la deuxième page
        document.add_page_break()
        second_page = document.add_paragraph()
        second_page.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = second_page.add_run("FEEDER:")
        run.bold = True
        run.underline = True
        run.font.size = Pt(20)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(
            0x2F, 0x54, 0x96)  # Bleu sombre (#2F5496)
        second_page.add_run(f"\n\n\t{feeder}\n\n")
        run = second_page.add_run("GROUPEMENT:")
        run.bold = True
        run.underline = True
        run.font.size = Pt(20)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(
            0x2F, 0x54, 0x96)  # Bleu sombre (#2F5496)
        second_page.add_run(f"\n\n\tGROUPEMENT TRONCONS ENTRE {groupement}")
        run.font.size = Pt(20)
        run.font.name = "Times New Roman"
        # Ajouter la troisième page
        document.add_page_break()
        third_page = document.add_paragraph()
        third_page.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = third_page.add_run("\n")
        #run.add_picture(Config.ASSETS_ROOT + '/img/logo.png', width=Inches(7))

        run = third_page.add_run("\n\n")
        run = third_page.add_run(f"GROUPEMENT TRONCONS ENTRE {groupement}")
        run.underline = True
        run.font.size = Pt(16)
        for image_info in image_data:
            # Charger l'image à partir des données binaires
            image_data_bytes = base64.b64decode(image_info.data)
            image = PILIMAGE.open(BytesIO(image_data_bytes))

            # Récupérer les informations sur les défauts
            # Supposant que les défauts sont séparés par des barres obliques
            defects = image_info.type_defaut.split('/')

            for defect in defects:
                defect = defect.strip()
                # Accédez aux remarques et conseils depuis le fichier JSON
                remarque = "Remarque par défaut"
                conseil = "Conseil par défaut"
                for defaut, info in normes_conseils_data.items():
                    if defect in defaut:
                        remarque = info.get("I", "Remarque par défaut")
                        conseil = info.get("J", "Conseil par défaut")
                        break

                # Ajouter une page pour chaque défaut
                document.add_page_break()
                page = document.add_paragraph()
                page.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                # Créer le tableau pour stocker les données de l'image
                table = document.add_table(rows=5, cols=1)
                table.style = "Table Grid"
                table.columns[0].width = Pt(900)

                # Cellule pour les défauts
                cell = table.cell(0, 0)
                cell.text = f"Défauts : {defect}"
                cell.paragraphs[0].runs[0].bold = True

                # Reduce image dimensions
                max_image_size = (800, 600)
                try:
                    image.thumbnail(max_image_size, PILIMAGE.ANTIALIAS)
                except:
                    image.thumbnail(max_image_size)

                # Compress image to reduce size
                image = image.convert("RGB")
                image_bytes = BytesIO()
                image.save(image_bytes, format='JPEG', quality=95)
                image_bytes.seek(0)

                # Cellule pour l'image
                cell = table.cell(1, 0)
                cell.vertical_alignment = 1  # Alignement vertical centré
                cell.add_paragraph().add_run().add_picture(image_bytes, width=Inches(7))

                # Cellule pour la localisation
                cell = table.cell(2, 0)
                cell.vertical_alignment = 1  # Alignement vertical centré
                # Cellule pour la remarque
                cell = table.cell(3, 0)
                cell.vertical_alignment = 1
                cell.add_paragraph().add_run("Remarque: ").bold = True
                cell.add_paragraph(remarque)

                cell = table.cell(4, 0)
                cell.vertical_alignment = 1
                cell.add_paragraph().add_run("Conseil: ").bold = True
                cell.add_paragraph(conseil)

        # Ajouter le pied de page
        footer_text = f"RAPPORT DU {date}\t\t{nom_operateur}"
        sections[-1].footer.paragraphs[0].text = footer_text
        # En-tête
        header_text = "VOCASARA S.U.A.R.L"
        header_paragraph = sections[0].header.paragraphs[0]
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_run = header_paragraph.add_run(header_text)
        header_run.bold = True
        header_run.font.size = Pt(22)
        header_run.font.name = "Times New Roman"
        header_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Noir
        header_run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Logo
        #logo_path = Config.ASSETS_ROOT + '/img/logo.png'
        #header_paragraph.add_run().add_picture(logo_path, width=Inches(4)
        #                                       ).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Sauvegarder le document Word
        # Convert the date to a string with the format "YYYY-MM-DD"
        date = datetime.now().strftime("%Y-%m-%d")
        # document.save("rapport_du_feeder_{}_du_{}.docx".format(feeder,date))
        nom_du_fichier = ("Rapport_defauts_visible_du_feeder_{}_du_{}_par_{}".format(
            feeder, current_date, nom_operateur))
 
        # Enregistrer le fichier Word en mémoire
        word_file = io.BytesIO()
        document.save(word_file)
        word_file.seek(0)

        # Enregistrer le fichier dans la base de données
        document_word = DocumentRapportGenere(
            nom_operateur=nom_operateur,
            nom_du_rapport=nom_du_fichier,
            data=word_file.read(),  # Utilisez read() pour obtenir les données binaires
            type_de_fichier='word'
        )
        db.session.add(document_word)
        db.session.commit()
        success_message = "Le rapport a été généré avec succès."
        rapports = RapportGenere.query.all()
        generate_quantification_report()
        generate_resume_rapport()
        
    except Exception as e:
        logging.error(f"Une erreur s'est produite dans generate_doc : {str(e)}")
        error_message = f"Une erreur s'est produite : {str(e)}"


    return render_template('rapport/creer_un_rapport.html', rapports=rapports, success_message=success_message, error_message=error_message)


#@blueprint.route('/generate_quantification_report', methods=['GET', 'POST'])
def generate_quantification_report():
    success_message = None
    error_message = None
    rapports = []
    try:
        # Récupérer les données de la base de données
        current_date = datetime.now().strftime("%Y-%m-%d")
        last_image_data = ImageUploadVisible.query.order_by(
            ImageUploadVisible.upload_date.desc()).first()

        if not last_image_data:
            error_message = "Aucune donnée disponible dans la base de données."
            return render_template('rapport/creer_un_rapport.html')

        # Extraire les données nécessaires de la base de données
        feeder = last_image_data.feeder
        date = datetime.now().strftime("%Y-%m-%d")
        nom_operateur = last_image_data.nom_operateur
        image_data = ImageUploadVisible.query.filter_by(
            nom_operateur=nom_operateur).all()

        # Charger le modèle de feuille Quantification
        template_workbook = load_workbook("./apps/Exemple_rapport/Quantification_Statistique_exemple.xlsx")
        feuille_copy = template_workbook.active

        feuille_copy['D7'] = f"Noms : {nom_operateur}"
        feuille_copy['D8'] = f"Date : {date}"
        feuille_copy['D11'] = feeder

        # Iterate through image_data and count defects
        defect_count = {}
        for image_info in image_data:
            if image_info.type_defaut:
                defects = image_info.type_defaut.split('/')
                for defect in defects:
                    defect = defect.strip()
                    if defect and '/' not in defect:
                        if defect in defect_count:
                            defect_count[defect] += 1
                        else:
                            defect_count[defect] = 1


        # Write data to the Quantification sheet
        row_num = 12
        for defect, count in defect_count.items():
            feuille_copy.append(['', '', defect, count])

        # Total of defects in a green cell
        feuille_copy.append(['', '', 'Totale des défauts', sum(defect_count.values())])
        total_cell = feuille_copy.cell(row=feuille_copy.max_row, column=4)
        total_cell.font = Font(name="Calibri", size=12, color="00AA00")

        # Merge cell containing "DRS" with cells below
        drs_cell = feuille_copy['A12']
        feuille_copy.merge_cells(start_row=12, start_column=1, end_row=feuille_copy.max_row, end_column=1)
        drs_cell.alignment = Alignment(horizontal="center", vertical="center")

        # Add Pie Chart at the end
        pie = DoughnutChart()
        labels = Reference(feuille_copy, min_col=3, min_row=14, max_row=feuille_copy.max_row - 1)
        data = Reference(feuille_copy, min_col=4, min_row=11, max_row=feuille_copy.max_row - 1)
        pie.add_data(data, titles_from_data=True)
        pie.set_categories(labels)
        chart_cell = feuille_copy.cell(row=feuille_copy.max_row + 2, column=1)
        feuille_copy.add_chart(pie, chart_cell.coordinate)

    # Sauvegarder le rapport de quantification
        nom_du_fichier = f"Quantification_Rapport_du_feeder_{feeder}_du_{current_date}_par_{nom_operateur}"

        # Sauvegarder le fichier en tant que données binaires dans la base de données
        excel_file = io.BytesIO()
        template_workbook.save(excel_file)
        excel_file.seek(0)

        # Ajouter le rapport dans la base de données
        rapports = RapportGenere.query.all()
        document_quantification = DocumentRapportGenere(
            nom_operateur=nom_operateur,
            nom_du_rapport=nom_du_fichier,
            data=excel_file.read(),
            type_de_fichier='excel'
        )
        db.session.add(document_quantification)
        db.session.commit()
        
        success_message = "Le rapport quantification a été généré avec succès."
    except FileNotFoundError:
        error_message = "Le fichier de template n'a pas été trouvé."
    except Exception as e:
        error_message = f"Une erreur s'est produite : {str(e)}"
        logging.error(f"Une erreur s'est produite dans generate_quantification_report: {str(e)}")

    return render_template('/rapport/creer_un_rapport.html', rapports=rapports, success_message=success_message, error_message=error_message)


def generate_resume_rapport():
    success_message = None
    error_message = None
    rapports = []
    try:
        # Récupérer les données depuis la base de données ImageUploadVisible
        curent_date = datetime.now().strftime("%Y-%m-%d")

        # Récupérer les paramètres de la requête
        last_image_data = ImageUploadVisible.query.order_by(
            ImageUploadVisible.upload_date.desc()).first()

        # Charger les normes_conseils des défauts à partir du fichier JSON
        with open('./apps/phrase_normes_conseils/normes_conseils.json', encoding='utf-8') as f:
            normes_conseils_data = json.load(f)

        # Utiliser les valeurs récupérées de la base de données
        feeder = last_image_data.feeder
        # Convert the date to a string with the format "YYYY-MM-DD"
        date = datetime.now().strftime("%Y-%m-%d")
        nom_operateur = last_image_data.nom_operateur
        zone = last_image_data.zone
        image_data = ImageUploadVisible.query.filter_by(
            nom_operateur=nom_operateur).all()

        # Créer une copie du fichier
        wb_copy = load_workbook(
            './apps/Exemple_rapport/resume_rapport_visibles_exemple.xlsx')
        feuille_copy = wb_copy.active

        # Écrire dans les cellules spécifiques du fichier d'origine
        feuille_copy['A8'] = "Feeder : " + feeder
        feuille_copy['F3'] = "Date : " + date
        feuille_copy['G6'] = "Nom opérateur : " + nom_operateur
        feuille_copy['G8'] = "Zone : " + zone
        # Définir les styles de cellule
        font = Font(name='Calibri', size=18)
        alignment = Alignment(horizontal='center', vertical='center')
        border = Border(top=Side(border_style='thick'),
                        bottom=Side(border_style='thick'),
                        left=Side(border_style='thick'),
                        right=Side(border_style='thick'))

        # Générer des données pour chaque colonne
        row_num = 12
        for image_info in image_data:
            if image_info.type_defaut is not None and image_info.type_defaut.strip() != "":
                defauts = image_info.type_defaut.split("/")
                defauts = [defaut.strip() + '/' for defaut in defauts]
                
                colonne_I_values = []
                colonne_J_values = []

                for defaut in defauts:
                    if defaut in normes_conseils_data:
                        colonne_I_values.append(normes_conseils_data[defaut]['I'])
                        colonne_J_values.append(normes_conseils_data[defaut]['J'])

                feuille_copy[f'A{row_num}'] = image_info.upload_date  # Date/Heure
                feuille_copy[f'B{row_num}'] = feeder  # feeder
                feuille_copy[f'C{row_num}'] = image_info.troncon  # troncon
                # longeur (à remplacer par la vraie valeur)
                feuille_copy[f'D{row_num}'] = ""
                feuille_copy[f'E{row_num}'] = image_info.filename  # Nom de l'image
                latitude_nom = "Latitude"
                longitude_nom = "Longitude"
                feuille_copy[f'F{row_num}'] = f"{latitude_nom} {float(image_info.latitude):.8f}, {longitude_nom} {float(image_info.longitude):.8f}"
                # Défaut de l'image
                feuille_copy[f'G{row_num}'] = image_info.type_defaut
                # urgences (à remplacer par la vraie valeur)
                feuille_copy[f'H{row_num}'] = ""
                # Écriture des valeurs dans les colonnes I et J
                feuille_copy[f'I{row_num}'] = ', '.join(colonne_I_values)
                feuille_copy[f'J{row_num}'] = ', '.join(colonne_J_values)
                feuille_copy[f'K{row_num}'] = row_num - 11  # Compter de 1 à n

                # Appliquer les styles aux cellules
                for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K']:
                    cell = feuille_copy[f'{col}{row_num}']
                    cell.alignment = alignment
                    cell.font = font
                    cell.border = border

                # Définir la hauteur des lignes
                feuille_copy.row_dimensions[row_num].height = 138

                row_num += 1  # Incrémenter le numéro de ligne

        # Appliquer le style au titre "Urgences"
        cell_title = feuille_copy['H11']
        cell_title.fill = openpyxl.styles.PatternFill(
            fill_type='solid', fgColor='FF0000')  # Rouge
        # Texte blanc en gras
        cell_title.font = Font(color="FFFFFF", bold=True)
        nom_du_fichier = "Resume_des_rapport_visibles_du_feeder_{}_du_{}_par_{}".format(
            feeder, curent_date, nom_operateur)
        # Enregistrer le fichier Excel en mémoire
        excel_file = io.BytesIO()
        wb_copy.save(excel_file)
        excel_file.seek(0)
        rapports = RapportGenere.query.all()
        # Enregistrer le fichier dans la base de données
        document = DocumentRapportGenere(
            nom_operateur=nom_operateur,
            nom_du_rapport=nom_du_fichier,
            data=excel_file.read(),
            type_de_fichier='excel'
        )
        db.session.add(document)
        db.session.commit()
        success_message = "Le rapport a été généré avec succès."
    except Exception as e:
        # Handle exceptions or errors
        error_message = f"Une erreur s'est produite : {str(e)}"
        logging.error(f"Une erreur s'est produite dans generate_resume: {str(e)}")
    return render_template('/rapport/creer_un_rapport.html', rapports=rapports, success_message=success_message, error_message=error_message)



@blueprint.route('/mes_rapports')
@login_required
def mes_rapports():
    # Récupérer l'utilisateur connecté
    #user = current_user
    rapports = DocumentRapportGenere.query.all()
    return render_template('rapport/mes_rapports.html', rapports=rapports)


@blueprint.route('/telecharger_rapport/<int:rapport_id>')
@login_required
def telecharger_rapport(rapport_id):
    try:
        rapport = DocumentRapportGenere.query.get_or_404(rapport_id)
        
        if rapport.type_de_fichier == 'excel':
            # Si le type de fichier est 'excel', c'est un document Excel
            return send_file(
                io.BytesIO(rapport.data),
                as_attachment=True,
                download_name=f"{rapport.nom_du_rapport}.xlsx",
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'  # Mimetype pour Excel
            )
        elif rapport.type_de_fichier == 'word':
            # Si le type de fichier est 'word', c'est un document Word
            return send_file(
                io.BytesIO(rapport.data),
                as_attachment=True,
                download_name=f"{rapport.nom_du_rapport}.docx",
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'  # Mimetype pour Word
            )
        else:
            # Gérer d'autres types de fichiers si nécessaire
            error_message='Type de fichier non pris en charge'

    except Exception as e:
        error_message = f"Erreur lors du téléchargement du rapport : {str(e)}"
        rapports = DocumentRapportGenere.query.all()
        return render_template('rapport/mes_rapports.html', error_message=error_message, rapports=rapports)


#./////////////////////////Partie inspections.///////////////


#chatbot./////////////////pas de login required ici
import uuid
import secrets
from apps.home.chatbot import get_response
from flask_limiter import Limiter
import logging
from flask_cors import CORS


CORS(blueprint)  # Add this line to enable CORS for your blueprint
# Define your model for the remember table
class Remember(db.Model):
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    user_id = db.Column(db.String(255))
    phrase = db.Column(db.Text)

# Vérifier si le dossier existe, sinon le créer
log_folder = './log'
if not os.path.exists(log_folder):
    os.makedirs(log_folder)
    
# Configuration du logging
logging.basicConfig(filename='./log/crash.log', level=logging.ERROR)

# Configuration du logging
access_logger = logging.getLogger('access')
access_logger.setLevel(logging.INFO)
file_handler = logging.FileHandler('./log/access.log')
access_logger.addHandler(file_handler)

limiter = Limiter(blueprint, default_limits=["200 per day", "100 per hour"])

def generate_secret_key():
    return secrets.token_urlsafe()

@blueprint.route('/chatbot')
def chatbot():
    return render_template('Chatbot/chatbot.html')


def generate_session_id():
    return str(uuid.uuid4())

# Fonction pour récupérer l'identifiant de session et l'identifiant de l'utilisateur
def get_user_ids():
    session_id = request.cookies.get('session_id')
    user_id = request.cookies.get('user_id')
    if not session_id:
        session_id = generate_session_id()
    if not user_id:
        user_id = generate_session_id()
    return session_id, user_id

@blueprint.route('/predict', methods=['POST'])
@limiter.limit("20 per minute")
def predict():
    try:
        message = request.get_json().get('message')
        session_id, user_id = get_user_ids()
        response = get_response(session_id, user_id, message)
        message = {"answer": response}
        resp = jsonify(message)
        resp.set_cookie('session_id', session_id, httponly=True, secure=True, samesite='Strict')
        resp.set_cookie('user_id', user_id, httponly=True, secure=True, samesite='Strict')
        return resp
    except Exception as e:
        logging.exception('Une erreur s\'est produite :')
        return jsonify({'error': str(e)}), 500


@blueprint.route('/save-phrase', methods=['POST'])
def save_phrase():
    try:
        user_id = request.cookies.get('user_id')
        phrase = request.get_json().get('phrase')

        # Create a new Remember instance and add it to the session
        remember_instance = Remember(user_id=user_id, phrase=phrase)
        db.session.add(remember_instance)
        db.session.commit()

        return jsonify({'success': 'Phrase saved successfully.'})
    except Exception as e:
        logging.exception('Une erreur s\'est produite :')
        return jsonify({'error': 'Une erreur s\'est produite.'}), 500

#chatbot //////////////////////////// pas de login required ici

@blueprint.route('/chatbot_info')
@login_required
def chatbot_info():
    return render_template('/Chatbot/chatbot_infos.html')

@blueprint.route('/confidentialite')
@login_required
def confidentialite():
    return render_template('home/confidentialite.html')